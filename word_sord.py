from docx import Document
import re

# Load the uploaded document
input_path = "We.docx"
output_path = "نمونه سوالات نانوایی مرتب شده.docx"

# Open the document
doc = Document(input_path)

# Prepare a new document for the sorted questions
new_doc = Document()

# Patterns for identifying questions and answers
question_pattern = re.compile(r"^[0-9]*\s?[-]?.*\?$")
answer_pattern = re.compile(r"^[الف-ی]:|پاسخ")

question_number = 1

# Process paragraphs
for para in doc.paragraphs:
    text = para.text.strip()
    if text:
        if question_pattern.match(text):  # Match questions
            new_doc.add_paragraph(f"{question_number}. {text.lstrip('').strip()}")
            question_number += 1
        elif answer_pattern.match(text):  # Match answers
            new_doc.add_paragraph(f"   - {text.strip()}")

# Save the formatted document
new_doc.save(output_path)
