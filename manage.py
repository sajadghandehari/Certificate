import sys
import time
from PyQt5.QtGui import *
from PyQt5.QtCore import *
from datetime import datetime
from PyQt5.QtWidgets import *
from PyQt5.uic import loadUiType
from front import style
import pandas as pd
import shutil
import os
from PyQt5.QtCore import QDate
from khayyam import JalaliDate
from PyQt5.QtCore import Qt
from docx import Document
from docx.shared import Pt  
from docx.oxml.ns import qn 
from docx.oxml import OxmlElement
import comtypes.client
from PyPDF2 import PdfMerger
import re  # برای بررسی اعداد
import random
import traceback



mainWindow, _ = loadUiType('main_page.ui')
errorWindow, _ = loadUiType('error_message.ui')
errorWindow_2, _ = loadUiType('error_message_2.ui')


class ErrorMessage(QDialog, errorWindow):
    def __init__(self, message):
        QDialog.__init__(self)


        self.setupUi(self)
        self.setWindowTitle("Error")
        self.Handel_Buttons()
        self.explain_error.setText(message)

    def Handel_Buttons(self):

        self.close_button.clicked.connect(self.close_window)
 

    def close_window(self):
        self.close() 

class ErrorMessage2(QDialog, errorWindow_2):
    def __init__(self):
        QDialog.__init__(self)

        self.setupUi(self)
        self.setWindowTitle("Error")
        # self.setWindowIcon(QIcon('icon/python4.webp'))
        self.Handel_Buttons()
 
    def Handel_Buttons(self):

        self.close_button.clicked.connect(self.close_window)
 

    def close_window(self):
        self.close() 


class Mainwindow(QMainWindow, mainWindow):

    def __init__(self):
        QMainWindow.__init__(self)
        self.setupUi(self)
        self.progressBar.setVisible(False)


        self.setWindowTitle("Certificate Issuance App")
        # self.setWindowIcon(QIcon('icon/python4.webp'))
        min_date = QDate(1200, 1, 1)
        self.dateEdit.setMinimumDate(min_date)
        self.Handel_Buttons()
 
        # self.submit_button.clicked.connect(self.submit)
        self.calasic_theme()
        self.theme = 'clasic'

        # self.update_data()
        self.comboBox.setCurrentIndex(0) 
        # self.calender(self.comboBox.currentText())
        # self.comboBox.currentTextChanged.connect(self.calender)
        self.insert_file.clicked.connect(self.open_file_dialog)
        today_date = JalaliDate.today()
        qdate = QDate(today_date.year, today_date.month, today_date.day)
        self.dateEdit.setMinimumDate(QDate(1000, 1, 1))
        self.dateEdit_2.setMinimumDate(QDate(1000, 1, 1))
        self.dateEdit_3.setMinimumDate(QDate(1000, 1, 1))
        self.dateEdit.setDate(qdate)
        self.dateEdit_2.setDate(qdate)
        self.dateEdit_3.setDate(qdate)


        def clear_only_files(folder_path):
            if os.path.exists(folder_path):
                for file in os.listdir(folder_path):
                    file_path = os.path.join(folder_path, file)
                    if os.path.isfile(file_path):
                        os.unlink(file_path)  # حذف فایل

        clear_only_files('database')



    def button_clicked(self, row):

        self.df.drop(row, inplace=True)
        # Save the modified DataFrame back to Excel
        self.df.to_excel('MOCK_DATA.xlsx', index=False)
        self.update_data()


    def Handel_Buttons(self):

        self.submit_button.clicked.connect(lambda: self.certificate_issuance(certificate_type='issuance'))
        self.submit_button_2.clicked.connect(lambda: self.certificate_issuance(certificate_type='Preview'))
        # self.darktheme.clicked.connect(self.dark_theme)
        
        
    def clear_folder(self, folder_path, type='pdf'):
        # بررسی وجود پوشه
        if os.path.exists(folder_path) and os.path.isdir(folder_path):
            for filename in os.listdir(folder_path):
                file_path = os.path.join(folder_path, filename)
                try:
                    # حذف فایل‌های PDF
                    if type == 'pdf':
                        if os.path.isfile(file_path) and file_path.lower().endswith('.pdf'):
                            os.remove(file_path)
                            print(f"فایل {file_path} حذف شد.")
                    else:
                        os.remove(file_path)
                except Exception as e:
                    print(f"خطا در حذف {file_path}: {e}")
                    print("An error occurred:")
                    print(e)
                    print("Traceback details:")
                    traceback.print_exc()
                    error_window = ErrorMessage2()
                    error_window.exec_()
                    break
                    
        else:
            print("مسیر پوشه معتبر نیست.")


    def certificate_issuance(self,certificate_type):

        print(certificate_type)

        shomare_govahi = self.lineEdit_1.text()
        moh = self.lineEdit_2.text()
        sal = self.lineEdit_3.text()
        selected_item = self.comboBox.currentText()
        name_dore = selected_item
        print(f"آیتم انتخاب‌شده: {selected_item}")
        ostan = self.lineEdit_5.text()
        sahrestan = self.lineEdit_6.text()
        number_font_size = int(self.lineEdit_7.text())
        text_font_size = int(self.lineEdit_8.text())
        text_font = 'IranNastaliq'
        number_font = 'B Nazanin'
        tarikh = self.dateEdit.date().toString('yyyy/MM/dd')
        tarikh_shro_dore = self.dateEdit_2.date().toString('yyyy/MM/dd')
        tarikh_payan_dore = self.dateEdit_3.date().toString('yyyy/MM/dd')


        def resource_path(relative_path):
            """برای دسترسی به فایل‌ها در محیط exe"""
            if hasattr(sys, '_MEIPASS'):
                return os.path.join(sys._MEIPASS, relative_path)
            return os.path.join(os.path.abspath("."), relative_path)
        excel_path = resource_path("database/data.xlsx")
        Preview_folder_path = resource_path("Preview")
        certificate_folder_path = resource_path("Certificates/Process")

        random_number = random.randint(1, 1000)
        preview_file_path_word = os.path.join(Preview_folder_path, f"certificate_{random_number}.docx")
        preview_file_path_pdf = os.path.join(Preview_folder_path, f"certificate_{random_number}.pdf")


        # folder_path = resource_path('Preview')
        if not os.path.exists(Preview_folder_path):
            os.makedirs(Preview_folder_path)
        # if not os.path.exists(excel_path):
        #     os.makedirs(excel_path)
        if not os.path.exists(certificate_folder_path):
            os.makedirs(certificate_folder_path)
        # print(folder_path)

        try:
            self.clear_folder(Preview_folder_path, 'all')
        except:
            print(174)
            error_window = ErrorMessage2()
            error_window.exec_()


        try:
            exel_path = resource_path('database\data.xlsx')
            # exel_path = 'database\data.xlsx'
            data = pd.read_excel(exel_path)  
            if name_dore == 'مهارت عمومی چهارگانه':
                template_path = 'template_1.docx'  
            else:
                template_path = 'template_2.docx'  
            self.progressBar.setVisible(True)


            for index, person in data.iterrows():

                if index == 1 and certificate_type == 'Preview':

                    # folder_path = "Preview"
                    if os.path.exists(preview_file_path_pdf):
                        print('exist file')
                        try:
                            self.clear_folder(folder_path)

                        except:
                            print(174)
                            error_window = ErrorMessage2()
                            error_window.exec_()
                    break
                
                main_data = {"ماه": moh, "سال": sal, "شهرستان": sahrestan, "استان": ostan, "شماره": shomare_govahi, "تاریخ": tarikh, "نام دوره": name_dore, 'تاریخ شروع': tarikh_shro_dore, 'تاریخ پایان':tarikh_payan_dore}
                exel_data = {"نام و نام خانوادگی": person["نام و نام خانوادگی"], "نام پدر": person["نام پدر"], "کد ملی": person["کد ملی"], "تاریخ تولد": person["تاریخ تولد"], "پایه خدمتی": person["پایه خدمتی"]}


                main_data.update(exel_data)
                doc = Document(template_path)  # Load the template file for each person
                
                
                # Replace placeholders with actual values
                
                default_font = text_font
                normal_font = number_font

                count = 0
                for paragraph in doc.paragraphs:

                    for key, value in main_data.items():
                        if f"{{{{{key}}}}}" in paragraph.text:
                            # جایگزینی متن در پاراگراف
                            paragraph.text = paragraph.text.replace(f"{{{{{key}}}}}", str(value))
                            
                            if count < 2  :

                                print(f"{key} Value:", value)
                                count += 1
                                for run in paragraph.runs:
                                    run.font.name = default_font  # تنظیم فونت
                                    run.font.size = Pt(text_font_size-2)  # تنظیم اندازه فونت
                                    run.bold = True 
                                    
                                    # پشتیبانی از زبان فارسی
                                    r = run._element
                                    rPr = r.get_or_add_rPr()
                                    rFonts = OxmlElement('w:rFonts')
                                    rFonts.set(qn('w:ascii'), default_font)
                                    rFonts.set(qn('w:hAnsi'), default_font)
                                    rFonts.set(qn('w:cs'), default_font)
                                    rPr.append(rFonts)
                                continue


                            for run in paragraph.runs:
                                words = run.text.split()  # جدا کردن کلمات در متن run

                                run.text = ""  # پاک کردن متن اصلی run

                                for word in words:
                                    new_run = paragraph.add_run(word + " ")
                                    
                                    # ابتدا تمام تنظیمات پیشفرض را غیرفعال کنید
                                    new_run.font.name = default_font  # پیشفرض
                                    new_run.font.size = Pt(text_font_size)
                                    new_run.bold = True

                                    # سپس شرطهای خاص را بررسی کنید
                                    if word.strip() == str(value).strip():
                                        print(f"{key} Value:", word)
                                        new_run.font.name = normal_font
                                        new_run.font.size = Pt(40)
                                        current_font = normal_font  # برای استفاده در XML
                                    elif re.match(r"^\d+(\.\d+)?$", word):
                                        new_run.font.name = normal_font
                                        new_run.font.size = Pt(number_font_size)
                                        current_font = normal_font
                                    elif re.match(r"^\d+/\d+/\d+$", word):
                                        new_run.font.name = normal_font
                                        new_run.font.size = Pt(number_font_size)
                                        current_font = normal_font
                                    else:
                                        current_font = default_font

                                    # تنظیمات XML با فونت صحیح
                                    r = new_run._element
                                    rPr = r.get_or_add_rPr()
                                    rFonts = OxmlElement('w:rFonts')
                                    rFonts.set(qn('w:ascii'), current_font)
                                    rFonts.set(qn('w:hAnsi'), current_font)
                                    rFonts.set(qn('w:cs'), current_font)
                                    rPr.append(rFonts)



                # ذخیره فایل
                if certificate_type == 'Preview':
                    # Preview_file_path_word = os.path.abspath(".\\Preview\\certificate_1.docx")
                    doc.save(preview_file_path_word)
                else:
                    certificate__file_path_word = os.path.join(certificate_folder_path, f"certificate_{index+1}.docx")
                    doc.save(certificate__file_path_word)

            print("گواهی‌ها با موفقیت ایجاد شدند.")




            # مسیر پوشه‌ای که فایل‌های ورد در آن قرار دارند
            if certificate_type == 'Preview':
                # folder_path = os.path.abspath(".\\Preview")
                folder_path = Preview_folder_path

                input_files = sorted([f for f in os.listdir(folder_path) if f.startswith('certificate') and f.endswith('.docx')])
                self.progressBar.setMaximum(10)
                self.progressBar.setValue(8)

            else:
                # folder_path = os.path.abspath(".\\Certificates\\Process")
                folder_path = certificate_folder_path

                input_files = sorted([f for f in os.listdir(folder_path) if f.startswith('certificate') and f.endswith('.docx')])
                self.progressBar.setMaximum(len(input_files))
                self.progressBar.setValue(0)

            # فهرست فایل‌های ورد در پوشه، بر اساس نام فایل‌ها

            # تبدیل فایل ورد به PDF
            def convert_to_pdf(docx_path, pdf_path, certificate_type):
                word = comtypes.client.CreateObject('Word.Application')
                doc = word.Documents.Open(docx_path)
                doc.SaveAs(pdf_path, FileFormat=17)  # فرمت PDF
                doc.Close()
                word.Quit()
                if certificate_type == 'issuance':
                    os.remove(docx_path)  # حذف فایل ورد پس از تبدیل

            # تبدیل همه فایل‌های ورد به PDF
            count = 1

            try:
                for file in input_files:

                    file_path = os.path.join(folder_path, file)
                    pdf_output = os.path.join(folder_path, file.replace('.docx', '.pdf'))
                    if certificate_type == 'issuance':
                        convert_to_pdf(file_path, pdf_output,'issuance')
                    else:
                        convert_to_pdf(file_path, pdf_output,'Preview')

                    self.progressBar.setValue(count)
                    count += 1
            except Exception as e:
                print("An error occurred:")
                print(e)
                print("Traceback details:")
                traceback.print_exc()
                error_window = ErrorMessage2()
                error_window.exec_()
            # فهرست فایل‌های PDF در پوشه
            input_pdf_files = sorted([f for f in os.listdir(folder_path) if f.startswith('certificate') and f.endswith('.pdf')])

            if certificate_type == 'issuance':

                # ترکیب فایل‌های PDF
                pdf_merger = PdfMerger()
                for file in input_pdf_files:
                    file_path = os.path.join(folder_path, file)
                    pdf_merger.append(file_path)

                # ذخیره کردن فایل ترکیب‌شده به عنوان یک PDF جدید
                try:
                    output_pdf = os.path.join('Certificates', f'دوره_{moh}_{sal}_{name_dore}.pdf')
                    pdf_merger.write(output_pdf)
                except Exception as e:

                    print("An error occurred:")
                    print(e)
                    print("Traceback details:")
                    traceback.print_exc()
                    error_window = ErrorMessage2()
                    error_window.exec_()

                # بستن شیء PdfMerger
                pdf_merger.close()

                # حذف فایل‌های PDF پس از ترکیب
                for file in input_pdf_files:
                    file_path = os.path.join(folder_path, file)
                    os.remove(file_path)  # حذف فایل PDF

            # print(f"تمام فایل‌های ورد به PDF تبدیل شدند و فایل‌های PDF حذف شدند.")
            else:
                try:
                    # Preview_file_path_pdf = os.path.abspath(".\\Preview\\certificate_1.pdf")
                    os.startfile(preview_file_path_pdf)
                except Exception as e:
                    print("An error occurred:")
                    print(e)
                    print("Traceback details:")
                    traceback.print_exc()
                    error_window = ErrorMessage2()
                    error_window.exec_()

            self.progressBar.setVisible(False)

        except Exception as e:
            print("An error occurred:")
            print(e)
            print("Traceback details:")
            traceback.print_exc()

            self.show_error_message('لطفا فایل اکسل را وارد کنید !')


    def Preview_certificate(self):
        print('Preview_certificate')
 

    def open_file_dialog(self):
        file_dialog = QFileDialog()
        file_path, _ = file_dialog.getOpenFileName(self, "Open File")



        if file_path.endswith(('.xlsx')):
            
            def resource_path(relative_path):
                """برای دسترسی به فایل‌ها در محیط exe"""
                if hasattr(sys, '_MEIPASS'):
                    return os.path.join(sys._MEIPASS, relative_path)
                return os.path.join(os.path.abspath("."), relative_path)
            

            destination_folder = resource_path("database")
            os.makedirs(destination_folder, exist_ok=True)
            new_file_path = os.path.join(destination_folder, 'data.xlsx')

            try:
                shutil.copyfile(file_path, new_file_path)
            except PermissionError as e:
                print(f"PermissionError: {e}")
                if os.path.exists(new_file_path):
                    print("Check if the destination file is open or locked.")

            # shutil.copyfile(file_path, new_file_path)

        else:
            self.show_error_message('فایل انتخابی باید اکسل باشد !')

         
    def show_error_message(self, errorMessage):
        error_window = ErrorMessage(errorMessage)
        error_window.exec_()

    def calasic_theme(self):

        self.theme = 'clasic'


        self.submit_button.setStyleSheet(style.button('008012'))
        self.submit_button_2.setStyleSheet(style.button('008012'))



if __name__ == '__main__':
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon('icon/certificate.png'))
    window = Mainwindow()
    window.setFixedSize(900, 600)
    window.show()
    sys.exit(app.exec_())